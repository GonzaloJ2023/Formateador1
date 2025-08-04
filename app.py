# Importaciones necesarias para el servidor web y el procesamiento de Word.
import os
import io
import base64
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from lxml import etree

# --------------------------------------------------------------------------
# Configuración del Servidor Flask
# --------------------------------------------------------------------------

# Crea la instancia de la aplicación Flask.
# CORRECCIÓN: Usamos 'app' de forma consistente para toda la aplicación.
app = Flask(__name__)
# Habilita CORS para permitir que el frontend acceda a este servidor.
CORS(app)

# --------------------------------------------------------------------------
# Funciones Auxiliares para el Procesamiento del Documento
# --------------------------------------------------------------------------

def clean_and_process_document(document_bytes, text_to_remove=None):
    """
    Procesa un documento de Word desde datos binarios, eliminando párrafos
    que contengan un texto específico, y devuelve el documento modificado
    en un objeto BytesIO.
    """
    try:
        # Carga el documento de Word desde un buffer de memoria.
        documento = Document(io.BytesIO(document_bytes))

        parrafos_a_eliminar = []
        for parrafo in documento.paragraphs:
            parrafo_texto_limpio = parrafo.text.strip()
            
            # Revisa si el párrafo está vacío o si contiene el texto a remover
            es_parrafo_vacio = not parrafo_texto_limpio
            contiene_texto_a_remover = False
            if text_to_remove:
                contiene_texto_a_remover = text_to_remove.lower() in parrafo_texto_limpio.lower()

            if es_parrafo_vacio or contiene_texto_a_remover:
                parrafos_a_eliminar.append(parrafo)

        # Elimina los párrafos marcados del documento.
        for parrafo in parrafos_a_eliminar:
            p = parrafo._element
            p.getparent().remove(p)

        # Guarda el documento modificado en un buffer de memoria.
        modificado_bytes = io.BytesIO()
        documento.save(modificado_bytes)
        modificado_bytes.seek(0)
        return modificado_bytes

    except Exception as e:
        print(f"Error en el procesamiento del documento: {e}")
        return None

def get_html_preview(document_bytes):
    """
    Genera una vista previa en HTML del documento de Word, extrayendo el texto
    y manteniendo los saltos de línea.
    """
    try:
        documento = Document(io.BytesIO(document_bytes))
        html_content = ""
        for parrafo in documento.paragraphs:
            # Reemplaza saltos de línea con etiquetas <br> y envuelve en <p>
            text_with_breaks = parrafo.text.replace('\n', '<br/>')
            html_content += f"<p>{text_with_breaks}</p>"
        return html_content
    except Exception as e:
        return f"Error al generar la vista previa: {e}"

# --------------------------------------------------------------------------
# Ruta de la API
# --------------------------------------------------------------------------

@app.route('/process-document', methods=['POST'])
def process_document():
    """
    Maneja la petición POST desde el frontend, procesa el documento y devuelve
    una vista previa en HTML y el documento en Base64.
    """
    # Verifica que el archivo exista en la solicitud.
    if 'file' not in request.files:
        return jsonify({'error': 'No se encontró el archivo en la solicitud'}), 400

    file = request.files['file']
    text_to_remove = request.form.get('format_text', '')

    # Valida el archivo subido.
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo inválido'}), 400
    if not file.filename.endswith('.docx'):
        return jsonify({'error': 'Tipo de archivo no soportado. Por favor, sube un archivo .docx'}), 400

    try:
        # Lee el contenido del archivo en memoria.
        document_bytes = file.read()
        
        # Procesa el documento.
        processed_document_bytes = clean_and_process_document(document_bytes, text_to_remove)

        if processed_document_bytes:
            # Genera la vista previa en HTML del documento procesado.
            html_preview = get_html_preview(processed_document_bytes.getvalue())
            
            # Codifica el documento procesado en Base64 para el frontend.
            docx_base64 = base64.b64encode(processed_document_bytes.getvalue()).decode('utf-8')

            # Devuelve la respuesta en formato JSON.
            return jsonify({
                'html_content': html_preview,
                'docx_base64': docx_base64
            })
        else:
            return jsonify({'error': 'Error interno al procesar el documento.'}), 500

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# --------------------------------------------------------------------------
# Punto de entrada del script
# --------------------------------------------------------------------------

if __name__ == '__main__':
    # Ejecuta la aplicación en modo de desarrollo local.
    app.run(debug=True, host='0.0.0.0', port=os.environ.get("PORT", 5000))
